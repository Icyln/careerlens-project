from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document

try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover
    fitz = None

SUPPORTED_EXTENSIONS = {'.pdf', '.docx'}


def clean_text(text: str) -> str:
    text = text.replace('\x00', ' ')
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def extract_text_from_resume(file_path: str | Path) -> tuple[str, dict[str, Any]]:
    path = Path(file_path)
    extension = path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        return '', {
            'file_type': extension.replace('.', ''),
            'supported': False,
            'warnings': ['Unsupported file type. Please upload a PDF or DOCX resume.'],
        }

    if extension == '.pdf':
        return extract_pdf(path)
    return extract_docx(path)


def extract_pdf(path: Path) -> tuple[str, dict[str, Any]]:
    metadata: dict[str, Any] = {
        'file_type': 'pdf',
        'supported': True,
        'page_count': 0,
        'image_count': 0,
        'drawing_count': 0,
        'table_count': 0,
        'text_character_count': 0,
        'parser_quality': 'good',
        'warnings': [],
    }

    if fitz is None:
        metadata['parser_quality'] = 'failed'
        metadata['warnings'].append('PyMuPDF is not installed, so PDF text could not be extracted.')
        return '', metadata

    try:
        document = fitz.open(path)
        metadata['page_count'] = document.page_count
        page_texts: list[str] = []
        for page in document:
            page_texts.append(page.get_text('text') or '')
            metadata['image_count'] += len(page.get_images(full=True))
            try:
                metadata['drawing_count'] += len(page.get_drawings())
            except Exception:
                pass
        text = clean_text('\n'.join(page_texts))
        metadata['text_character_count'] = len(text)
        if len(text) < 250:
            metadata['parser_quality'] = 'poor'
            metadata['warnings'].append(
                'Very little selectable text was extracted. The resume may be scanned, image-heavy, or built with hidden text boxes.'
            )
        elif metadata['image_count'] > max(3, metadata['page_count'] * 2):
            metadata['parser_quality'] = 'medium'
            metadata['warnings'].append(
                'The PDF contains many images. Some applicant tracking systems may miss text embedded inside images, charts, or graphics.'
            )
        elif metadata['drawing_count'] > max(12, metadata['page_count'] * 8):
            metadata['parser_quality'] = 'medium'
            metadata['warnings'].append(
                'The PDF has many vector drawings. Keep graphics simple so ATS parsers can read the resume reliably.'
            )
        document.close()
        return text, metadata
    except Exception as exc:
        metadata['parser_quality'] = 'failed'
        metadata['warnings'].append(f'Could not parse PDF: {exc}')
        return '', metadata


def extract_docx(path: Path) -> tuple[str, dict[str, Any]]:
    metadata: dict[str, Any] = {
        'file_type': 'docx',
        'supported': True,
        'page_count': None,
        'image_count': 0,
        'drawing_count': 0,
        'table_count': 0,
        'text_character_count': 0,
        'parser_quality': 'good',
        'warnings': [],
    }

    try:
        document = Document(path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_text: list[str] = []
        for table in document.tables:
            metadata['table_count'] += 1
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_text.append(' | '.join(cells))
        metadata['image_count'] = len(document.inline_shapes)
        text = clean_text('\n'.join(paragraphs + table_text))
        metadata['text_character_count'] = len(text)
        if len(text) < 250:
            metadata['parser_quality'] = 'poor'
            metadata['warnings'].append(
                'Very little text was extracted from the DOCX file. Avoid text boxes, images, or heavily designed layouts.'
            )
        elif metadata['table_count'] > 5:
            metadata['parser_quality'] = 'medium'
            metadata['warnings'].append(
                'The DOCX uses many tables. Tables can be readable, but simpler one-column layouts are safer for ATS parsing.'
            )
        elif metadata['image_count'] > 2:
            metadata['parser_quality'] = 'medium'
            metadata['warnings'].append(
                'The DOCX includes images. Do not put important text inside images.'
            )
        return text, metadata
    except Exception as exc:
        metadata['parser_quality'] = 'failed'
        metadata['warnings'].append(f'Could not parse DOCX: {exc}')
        return '', metadata
